"""
Document Parser for RAG Ingestion Pipeline
Extracts text from various file formats: TXT, MD, PDF, DOCX
"""

from pathlib import Path
from typing import Optional
from dataclasses import dataclass

from PyPDF2 import PdfReader
from docx import Document


@dataclass
class ParsedDocument:
    """Result of parsing a document"""
    text: str
    page_count: int
    file_type: str
    original_filename: str


class DocumentParser:
    """
    Extracts text from various document formats.
    
    Supported formats: .txt, .md, .pdf, .docx
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.pdf', '.docx'}
    
    def supports(self, file_path: Path) -> bool:
        """Check if this file type is supported"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: Path) -> Optional[ParsedDocument]:
        """
        Parse a document and extract text.
        
        Args:
            file_path: Path to the file
            
        Returns:
            ParsedDocument with extracted text, or None if unsupported/failed
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if not self.supports(file_path):
            return None
        
        if extension in {'.txt', '.md'}:
            return self._parse_text(file_path)
        elif extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        
        return None
    
    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text or markdown files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return ParsedDocument(
            text=text,
            page_count=1,
            file_type=file_path.suffix.lower().lstrip('.'),
            original_filename=file_path.name
        )
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF files using PyPDF2"""
        reader = PdfReader(file_path)
        pages_text = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)
        
        return ParsedDocument(
            text="\n\n".join(pages_text),
            page_count=len(reader.pages),
            file_type="pdf",
            original_filename=file_path.name
        )
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX files using python-docx"""
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        return ParsedDocument(
            text="\n\n".join(paragraphs),
            page_count=1,  # DOCX doesn't have page concept in same way
            file_type="docx",
            original_filename=file_path.name
        )


# Singleton instance
_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """Get the global DocumentParser instance"""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
